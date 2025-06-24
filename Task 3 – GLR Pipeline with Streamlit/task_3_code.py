import streamlit as st
from docx import Document
import fitz  # PyMuPDF
import requests
import io
import json
import re

# -----------------------------
# Extract text from PDFs
# -----------------------------
def extract_text_from_pdfs(pdf_files):
    all_text = ""
    for pdf_file in pdf_files:
        pdf_bytes = pdf_file.read()
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in pdf_doc:
            all_text += page.get_text()
        pdf_doc.close()
        pdf_file.seek(0)
    return all_text

# -----------------------------
# Extract [FIELD] tags from DOCX
# -----------------------------
def get_template_fields(doc):
    fields = set()
    for para in doc.paragraphs:
        matches = re.findall(r"\[([A-Z0-9_]+)\]", para.text)
        for match in matches:
            fields.add(match)
    return list(fields)

# -----------------------------
# Replace placeholders in template
# -----------------------------
def fill_template(doc, key_values):
    for para in doc.paragraphs:
        for key, val in key_values.items():
            placeholder = f"[{key}]"
            safe_val = "" if val is None else str(val)
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, safe_val)
    return doc

# -----------------------------
# LLM Key-Value Pair & Narrative Generation
# -----------------------------
def get_key_value_pairs_and_narrative(template_fields, extracted_text):
    prompt = f"""
You are an expert insurance assistant.

Your task is two-fold:
1. Extract accurate values for the following placeholder fields from the inspection report below:
{template_fields}

2. Then generate a detailed full narrative insurance inspection report using the extracted values.

Return the result as valid JSON like this format:
{{
  "key_values": {{ "FIELD_NAME": "Value", ... }},
  "narrative_report": "Full report goes here..."
}}

Inspection Report:
{extracted_text}
"""

    API_URL = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": "Bearer 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",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek/deepseek-r1-0528-qwen3-8b:free",
        "messages": [
            {"role": "system", "content": "You extract structured insurance data and generate reports."},
            {"role": "user", "content": prompt}
        ]
    }

    try:
        response = requests.post(API_URL, headers=headers, json=data, timeout=180)
        response.raise_for_status()
        raw_output = response.json()['choices'][0]['message']['content']
        json_match = re.search(r'\{[\s\S]*\}', raw_output)
        json_text = json_match.group(0) if json_match else ""
        parsed = json.loads(json_text)
        return parsed.get("key_values", {}), parsed.get("narrative_report", ""), json_text
    except Exception as e:
        st.error(f"LLM or JSON parsing failed: {e}")
        return {}, "", ""

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="📄 AI Insurance Report Generator", layout="wide")
st.title("📄 AI Insurance Report Generator")

template_file = st.file_uploader("Upload insurance template (.docx)", type=["docx"])
pdf_files = st.file_uploader("Upload inspection report PDFs (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("🔄 Generate Report") and template_file and pdf_files:
    with st.spinner("Extracting content, generating key-values and narrative..."):
        template_doc = Document(template_file)
        extracted_text = extract_text_from_pdfs(pdf_files)
        template_fields = get_template_fields(template_doc)

        st.info("Extracted Fields from Template:")
        st.code(template_fields)

        key_values, narrative, raw_json = get_key_value_pairs_and_narrative(template_fields, extracted_text)

        if not key_values:
            st.error("❌ Failed to extract valid key-value pairs from model.")
        else:
            filled_doc = fill_template(template_doc, key_values)
            filled_doc.add_page_break()
            filled_doc.add_paragraph("Full Narrative Report:")
            filled_doc.add_paragraph(narrative)

            output_buffer = io.BytesIO()
            filled_doc.save(output_buffer)
            output_buffer.seek(0)

            st.success("✅ Report Ready!")
            st.download_button(
                label="📥 Download Final Report (.docx)",
                data=output_buffer,
                file_name="complete_insurance_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.subheader("🗂️ Extracted Key-Value Pairs")
            st.json(key_values)
            st.subheader("📝 Full Narrative")
            st.text_area("Narrative Report", narrative, height=350)
else:
    st.info("📎 Please upload a DOCX template and at least one PDF report to get started.")
